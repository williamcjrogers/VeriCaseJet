"""Evidence text extraction helpers.

These helpers exist to keep the Evidence Repository UI functional even when
Apache Tika is temporarily unavailable.

We intentionally support lightweight, dependency-local extraction for:
- .docx (python-docx)
- .xlsx (openpyxl)
- text-like files (.txt/.csv/.json/.xml/.md/.log)

Binary legacy formats (.doc/.xls) still generally require Tika.
"""

from __future__ import annotations

from io import BytesIO
from urllib.parse import urlparse, urlunparse


_TEXTY_EXTS: set[str] = {"txt", "csv", "json", "xml", "html", "md", "log"}


def tika_url_candidates(preferred: str | None) -> list[str]:
    """Return a small list of plausible Tika base URLs.

    In docker-compose, the service is typically resolvable as `tika`.
    In the provided Kubernetes manifests, the service is `tika-service`.

    We try the configured value first, then swap between these known service
    names when applicable.
    """

    candidates: list[str] = []

    def add(url: str | None) -> None:
        if not url:
            return
        u = url.strip().rstrip("/")
        if not u:
            return
        if u not in candidates:
            candidates.append(u)

    add(preferred)

    if not candidates:
        add("http://tika:9998")

    base = candidates[0]
    parsed = urlparse(base)
    host = parsed.hostname

    def with_host(new_host: str) -> str:
        netloc = new_host
        if parsed.port:
            netloc = f"{new_host}:{parsed.port}"
        rebuilt = parsed._replace(netloc=netloc)
        return urlunparse(rebuilt).rstrip("/")

    if host == "tika":
        add(with_host("tika-service"))
    elif host == "tika-service":
        add(with_host("tika"))

    # As a last resort, ensure both common names appear.
    add("http://tika:9998")
    add("http://tika-service:9998")

    return candidates


def _decode_text(content: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(enc)
        except Exception:
            continue
    return content.decode("utf-8", errors="replace")


def extract_text_from_bytes(
    content: bytes,
    *,
    filename: str | None = None,
    mime_type: str | None = None,
    max_chars: int = 2_000_000,
) -> str:
    """Best-effort text extraction without external services."""

    name = (filename or "").strip()
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    mime = (mime_type or "").lower().strip()

    if mime.startswith("text/") or ext in _TEXTY_EXTS:
        return _decode_text(content)[:max_chars]

    # DOCX
    if (
        ext == "docx"
        or mime
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        try:
            from docx import Document  # type: ignore

            doc = Document(BytesIO(content))
            parts: list[str] = []

            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    parts.append(t)

            # Tables (common in construction/commercial docs)
            for table in doc.tables:
                for row in table.rows:
                    cells = [(c.text or "").strip() for c in row.cells]
                    line = "\t".join(c for c in cells if c)
                    if line.strip():
                        parts.append(line)

            text = "\n".join(parts)
            return text[:max_chars]
        except Exception:
            return ""

    # XLSX
    if (
        ext == "xlsx"
        or mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    ):
        try:
            from openpyxl import load_workbook  # type: ignore

            wb = load_workbook(
                filename=BytesIO(content), read_only=True, data_only=True
            )
            try:
                out: list[str] = []
                total_chars = 0

                for sheet_name in wb.sheetnames:
                    if total_chars >= max_chars:
                        break
                    ws = wb[sheet_name]
                    header = f"--- Sheet: {sheet_name} ---"
                    out.append(header)
                    total_chars += len(header) + 1

                    for row in ws.iter_rows(values_only=True):
                        if total_chars >= max_chars:
                            break
                        cells = []
                        for v in row:
                            if v is None:
                                cells.append("")
                            else:
                                cells.append(str(v))
                        line = "\t".join(cells).rstrip()
                        if not line.strip():
                            continue
                        out.append(line)
                        total_chars += len(line) + 1

                return "\n".join(out)[:max_chars]
            finally:
                try:
                    wb.close()
                except Exception:
                    pass
        except Exception:
            return ""

    return ""
