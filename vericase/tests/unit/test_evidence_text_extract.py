import os
import sys
import unittest
from io import BytesIO


TEST_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
if TEST_ROOT not in sys.path:
    sys.path.insert(0, TEST_ROOT)


from api.app.evidence.text_extract import extract_text_from_bytes, tika_url_candidates


class TestEvidenceTextExtract(unittest.TestCase):
    def test_tika_url_candidates_swaps_service_name(self):
        candidates = tika_url_candidates("http://tika:9998")
        self.assertIn("http://tika:9998", candidates)
        self.assertIn("http://tika-service:9998", candidates)

        candidates2 = tika_url_candidates("http://tika-service:9998")
        self.assertIn("http://tika-service:9998", candidates2)
        self.assertIn("http://tika:9998", candidates2)

    def test_extract_docx_paragraphs_and_tables(self):
        try:
            from docx import Document  # type: ignore
        except Exception as e:
            self.skipTest(f"python-docx not available: {e}")

        doc = Document()
        doc.add_paragraph("Hello from Word")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Left"
        table.cell(0, 1).text = "Right"

        buf = BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        text = extract_text_from_bytes(content, filename="sample.docx")
        self.assertIn("Hello from Word", text)
        # Table extraction uses tab-separated values.
        self.assertIn("Left\tRight", text)

    def test_extract_xlsx_cells(self):
        try:
            from openpyxl import Workbook  # type: ignore
        except Exception as e:
            self.skipTest(f"openpyxl not available: {e}")

        wb = Workbook()
        ws1 = wb.active
        if ws1 is None:
            self.fail("openpyxl Workbook.active unexpectedly returned None")
        ws1.title = "Sheet1"
        ws1["A1"].value = "Alpha"
        ws1["B2"].value = 123

        ws2 = wb.create_sheet("Second")
        ws2.append(["X", "Y"])

        buf = BytesIO()
        wb.save(buf)
        content = buf.getvalue()

        text = extract_text_from_bytes(content, filename="sample.xlsx")
        self.assertIn("--- Sheet: Sheet1 ---", text)
        self.assertIn("Alpha", text)
        self.assertIn("--- Sheet: Second ---", text)
        self.assertIn("X\tY", text)


if __name__ == "__main__":
    unittest.main()
